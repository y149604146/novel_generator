from datetime import datetime

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor

from novel_writer.utilities import TextUtility
from novel_writer.utilities.Logger import logger

import markdown

class WordGenerator:
    """word文档生成器"""

    def __init__(self,
                title: str,
                version: str = "V1.0",
                save_full_path: str = "output",
                file_ext: str = ".docx"
                ):
        """
        初始化报告生成器
        :param title: 文档主标题
        :param version: 文档版本,与当前实时日期一同构成报告副标题
        :param save_full_path: 文档保存的全路径（含路径、文件名）
        """
        self.doc = Document()
        self.title = title
        self.version = version
        self.save_full_path = save_full_path+file_ext
        # 获取当前系统日期时间（格式：年-月-日 时:分:秒）
        self.current_datetime = datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")
        # 生成实时副标题
        self.subtitle = f"编写日期：{self.current_datetime} | 版本：{self.version}"

        self.default_font = "宋体"
        self.title_font = "微软雅黑"

    def _set_chinese_font(self, run, font_name: str = "宋体", size: Pt = Pt(12)):
        """设置中文字体（处理中文显示问题）"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 适配中文字体
        run.font.size = size
        return run

    def add_title_section(self):
        """添加标题部分"""
        # 主标题
        title_para = self.doc.add_heading(self.title, level=0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_para.runs[0]
        self._set_chinese_font(title_run, "微软雅黑", Pt(20))
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 128)  # 深蓝色

        # 副标题
        sub_para = self.doc.add_heading(self.subtitle, level=1)
        sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sub_run = sub_para.runs[0]
        self._set_chinese_font(sub_run, "宋体", Pt(12))

        # 空行分隔
        #self.doc.add_paragraph()

    def add_paragraph_section(self, heading: str, heading_style: str, content: str):
        """
        添加段落部分
        :param heading: 段落标题（如"1. 评估目的"）
        :param content: 段落内容
        """
        # 段落标题
        if heading != "":
            if heading_style!="":
                # self.doc.add_paragraph(heading, style='Heading 2')
                self.doc.add_paragraph(heading, style=heading_style)
            # heading_para = self.doc.add_paragraph(heading)
            # heading_run = heading_para.runs[0]
            # self._set_chinese_font(heading_run, "宋体", Pt(12))
            # heading_run.font.bold = True
            # heading_run.font.color.rgb = RGBColor(0, 0, 108)

        content = TextUtility.remove_block(content, "<think>", "</think>")
        print(f"Removed <think></think> content. "
              f"content = {content}")
        self.__parse_markdown(content)
            # 段落内容（首行缩进）
        # content_para = self.doc.add_paragraph(content)
        # content_para.paragraph_format.first_line_indent = Inches(0.4)
        # content_run = content_para.runs[0]
        # self._set_chinese_font(content_run, "宋体", Pt(12))
        #
        # # 空行分隔
        # self.doc.add_paragraph()

    def add_list_section(self, heading: str, items: list):
        """
        添加列表部分
        :param heading: 列表标题
        :param items: 列表项（每个元素为元组：(标题, 内容)）
        """
        # 列表标题
        if heading!="":
            self.doc.add_paragraph(heading, style='Heading 2')

        # 生成列表
        for title, content in items:
            list_para = self.doc.add_paragraph()
            list_para.add_run(f"• {title}：").bold = True
            list_para.add_run(content)
            for run in list_para.runs:
                self._set_chinese_font(run, "宋体", Pt(12))
            list_para.paragraph_format.space_before = Pt(0)  # 段前间距
            list_para.paragraph_format.space_after = Pt(0)  # 段后间距

        # 空行分隔
        # self.doc.add_paragraph()

    def add_table_section(self, heading: str, headers: list, data: list):
        """
        添加表格部分
        :param heading: 表格标题
        :param headers: 表头列表（如["评估维度", "得分"]）
        :param data: 表格数据（二维列表，如[[维度1, 得分1], [维度2, 得分2]]）
        """
        # 表格标题
        self.doc.add_paragraph(heading, style='Heading 2')

        # 创建表格（行数=数据行数+1表头行）
        table = self.doc.add_table(rows=1 + len(data), cols=len(headers))
        table.style = 'Table Grid'

        # 填充表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        # 填充数据
        for row_idx, row_data in enumerate(data):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_value in enumerate(row_data):
                row_cells[col_idx].text = str(cell_value)

        # 调整第一列宽度
        if len(headers) > 0:
            for cell in table.columns[0].cells:
                cell.width = Inches(1.8)

        # 空行分隔
        # self.doc.add_paragraph()

    def add_image_section(self, heading: str, image_path: str, width: Inches = Inches(5)):
        """
        添加图片部分
        :param heading: 图片标题
        :param image_path: 图片路径
        :param width: 图片宽度
        """
        # 图片标题
        self.doc.add_paragraph(heading, style='Heading 2')

        # 插入图片（处理文件不存在的情况）
        try:
            self.doc.add_picture(image_path, width=width)
        except FileNotFoundError:
            self.doc.add_paragraph("（提示：图片文件未找到，无法显示）")

        # 空行分隔
        #self.doc.add_paragraph()

    def save_document(self):
        """保存文档"""
        self.doc.save(self.save_full_path)
        logger.info(f"文档已成功保存至：{self.save_full_path}")

    def __parse_markdown(self, md_content: str):
        """
        解析内容：支持Markdown格式和纯文本
        - 含Markdown标记的部分按格式解析
        - 纯文本部分按段落保留（不丢失内容）
        """
        # 处理空内容
        if not md_content.strip():
            # para = self.doc.add_paragraph("（内容为空）")
            # self._set_chinese_font(para.runs[0])
            return

        # 1. 将内容转换为HTML（Markdown语法会被解析，纯文本会保留为<p>标签）
        try:
            # 启用表格扩展，其他扩展按需添加
            html_content = markdown.markdown(md_content, extensions=['tables'])
        except Exception as e:
            # 解析失败时直接按纯文本处理
            html_content = f"<p>{md_content.replace('\n', '<br>')}</p>"

        soup = BeautifulSoup(html_content, 'html.parser')
        elements = soup.contents

        # 2. 遍历解析后的元素，转换为Word内容
        for element in elements:
            if not element.name:
                # 处理纯文本节点（非标签内容）
                text = str(element).strip()
                if text:  # 跳过空文本
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.first_line_indent = Inches(0.4)
                    self._set_chinese_font(para.runs[0])
                    # self.doc.add_paragraph()
                continue

            # 处理标题（---）
            if element.name == 'hr':
                # 在Word中添加一条横线作为分隔（通过段落+下划线模拟）
                hr_para = self.doc.add_paragraph()
                hr_run = hr_para.add_run(" " * 50)  # 用空格撑开宽度
                hr_run.font.underline = True  # 下划线模拟横线
                hr_run.font.size = Pt(1)  # 细线效果
                # self.doc.add_paragraph()  # 分隔线后空行

            # 处理标题（h1-h6）
            elif element.name.startswith('h'):
                level = int(element.name[1:])+1
                word_level = min(level, 6)
                text = element.get_text(strip=True)
                heading = self.doc.add_heading(text, level=word_level)
                run = heading.runs[0]
                self._set_chinese_font(run, self.title_font, Pt(17 - level * 2))
                run.font.bold = True


            # 处理段落（p标签，包括纯文本段落）
            elif element.name == 'p':
                para = self.doc.add_paragraph()
                para.paragraph_format.first_line_indent = Inches(0.4)
                # 处理段落内的粗体（strong）、斜体（em）和普通文本
                for child in element.contents:
                    if hasattr(child, 'name'):
                        if child.name == 'strong':
                            run = para.add_run(child.get_text())
                            run.font.bold = True
                        elif child.name == 'em':
                            run = para.add_run(child.get_text())
                            run.font.italic = True
                        elif child.name == 'br':  # 处理换行
                            # para.add_run('\n')
                            print("will not deal with <br>.")
                        else:
                            run = para.add_run(child.get_text())
                            self._set_chinese_font(run)
                    else:
                        # 普通文本节点
                        run = para.add_run(str(child).strip())
                        self._set_chinese_font(run)
                # self.doc.add_paragraph()

            # 处理无序列表（ul）
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    li_text = li.get_text()
                    if li_text:
                        if str(li_text).strip() == "":
                            print("li content is ''.")
                        else:
                            para = self.doc.add_paragraph(li.get_text(), style='List Bullet')
                            self._set_chinese_font(para.runs[0])
                #self.doc.add_paragraph()

            # 处理有序列表（ol）
            elif element.name == 'ol':
                for li in element.find_all('li'):
                    li_text = li.get_text()
                    if li_text:
                        if str(li_text).strip() == "":
                            print("li content is ''.")
                        else:
                            # para = self.doc.add_paragraph(li.get_text(), style='List Number')
                            para = self.doc.add_paragraph(li.get_text(), style='List Bullet')
                            self._set_chinese_font(para.runs[0])
                #self.doc.add_paragraph()

            # 处理表格（table）
            elif element.name == 'table':
                headers = [th.get_text() for th in element.find('tr').find_all('th')] if element.find('tr') else []
                rows = []
                for tr in element.find_all('tr')[1:]:
                    rows.append([td.get_text() for td in tr.find_all('td')])

                if headers:  # 确保表头存在
                    table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
                    table.style = 'Table Grid'
                    # 填充表头
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = header
                        self._set_chinese_font(hdr_cells[i].paragraphs[0].runs[0])
                        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    # 填充内容
                    for row_idx, row_data in enumerate(rows):
                        row_cells = table.rows[row_idx + 1].cells
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < len(row_cells):  # 避免索引越界
                                row_cells[col_idx].text = cell_text
                                self._set_chinese_font(row_cells[col_idx].paragraphs[0].runs[0])
                    # self.doc.add_paragraph()


